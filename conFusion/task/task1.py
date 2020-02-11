from docx import Document
from docx.shared import Inches
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPDF, renderPM


def svg2pngtoword(svgFilepath,pngFilepath,wordFilepath):
    document = Document()
    document.add_heading('Svg Image Inserted', 0)
    p=document.add_paragraph()
    r=p.add_run()
    drawing = svg2rlg(svgFilepath)
    renderPM.drawToFile(drawing,pngFilepath, fmt="PNG")
    r.add_picture(pngFilepath)
    document.save(wordFilepath)

if __name__=='__main__':
    svg2pngtoword("001-powder.svg","gs.png",'gs.docx')